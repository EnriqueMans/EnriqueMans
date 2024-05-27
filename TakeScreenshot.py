import pyautogui
from PIL import Image
from docx import Document
from docx.shared import Inches
import os

# Take a screenshot
screenshot = pyautogui.screenshot()

# Save the screenshot to a temporary file
temp_image_path = 'temp_screenshot.png'
screenshot.save(temp_image_path)

# Check if the Word document exists, if not create a new one
doc_path = 'screenshots.docx'
if os.path.exists(doc_path):
    doc = Document(doc_path)
else:
    doc = Document()

# Add the screenshot to the Word document
doc.add_picture(temp_image_path, width=Inches(5))

# Save the Word document
doc.save(doc_path)

# Clean up the temporary image file
os.remove(temp_image_path)

print(f"Screenshot saved to {doc_path}")
